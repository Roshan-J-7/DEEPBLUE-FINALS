"""
DeepBlue Project Dossier Generator
Generates a professional Word document (~20-22 pages) for hackathon/demo presentation.

Usage:
    python docs/generate_dossier.py

Output:
    docs/DeepBlue_Project_Dossier.docx

Image placeholders are marked with [IMAGE: description].
Drop your actual images into docs/images/ and update the paths in each section function.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

# ── Paths ────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent
IMAGES_DIR = BASE_DIR / "images"
OUTPUT_PATH = BASE_DIR / "DeepBlue_Project_Dossier.docx"

# ── Ensure dirs ──────────────────────────────────────────────────────────
IMAGES_DIR.mkdir(exist_ok=True)

# ── Style helpers ────────────────────────────────────────────────────────

def set_document_defaults(doc: Document):
    """Set default font, margins, and styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Heading styles
    for level, size, color in [
        ("Heading 1", 22, RGBColor(0x0D, 0x47, 0xA1)),  # Deep blue
        ("Heading 2", 16, RGBColor(0x1A, 0x1A, 0x1A)),
        ("Heading 3", 13, RGBColor(0x33, 0x33, 0x33)),
    ]:
        h = doc.styles[level]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_image_placeholder(doc, description, width=Inches(5.5)):
    """Insert an actual image if it exists, otherwise a visible placeholder."""
    # Convention: image filename derived from description
    safe_name = description.lower().replace(" ", "_").replace("/", "_")
    for ext in (".png", ".jpg", ".jpeg"):
        img_path = IMAGES_DIR / f"{safe_name}{ext}"
        if img_path.exists():
            doc.add_picture(str(img_path), width=width)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return

    # Placeholder paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ IMAGE: {description} ]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()  # spacing


def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SECTION FUNCTIONS — Edit content inside each one
# ══════════════════════════════════════════════════════════════════════════

def section_01_cover(doc):
    """Page 1 — Cover Page"""
    doc.add_paragraph("\n\n\n")  # top spacing

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DeepBlue")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    run.font.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("AI-Powered Rural Healthcare Assessment System")
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph("\n")
    add_image_placeholder(doc, "cover_illustration", width=Inches(4.0))
    doc.add_paragraph("\n")

    for line in [
        "Team DeepBlue",
        "Institution: [Your Institution]",
        "Event: [Hackathon / Competition Name]",
        "Year: 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)

    add_page_break(doc)


def section_02_overview(doc):
    """Page 2 — Project Overview"""
    add_heading(doc, "1. Project Overview")

    add_heading(doc, "1.1 Problem Statement", level=2)
    add_body(doc,
        "Access to quality healthcare remains a critical challenge in rural and underserved "
        "communities. Limited availability of trained medical professionals, delayed diagnosis, "
        "and the absence of early intervention systems contribute to preventable health "
        "deterioration. Patients in remote areas often travel long distances for basic health "
        "assessments, resulting in late-stage detection of treatable conditions."
    )

    add_heading(doc, "1.2 Proposed Solution", level=2)
    add_body(doc,
        "DeepBlue is an AI-powered healthcare assessment platform designed for rural deployment. "
        "The system combines an intelligent questionnaire engine, wearable sensor integration, "
        "voice-based AI interaction, and a centralized monitoring portal to deliver accessible, "
        "early-stage health assessments without requiring on-site medical expertise."
    )
    add_bullet(doc, "AI-driven health assessment through structured questionnaires")
    add_bullet(doc, "Kiosk-based deployment for community health centers")
    add_bullet(doc, "Wearable sensor data integration for vital sign monitoring")
    add_bullet(doc, "NGO and admin monitoring portal for disease tracking and coordination")

    add_page_break(doc)


def section_03_system_overview(doc):
    """Page 3 — System Overview"""
    add_heading(doc, "2. System Overview")

    add_heading(doc, "2.1 What is DeepBlue", level=2)
    add_body(doc,
        "DeepBlue is a modular healthcare platform that brings AI-assisted medical assessment "
        "to locations where traditional healthcare infrastructure is limited. It operates through "
        "multiple interfaces — mobile application, kiosk terminal, and voice-based AI caller — "
        "to maximize accessibility across diverse user populations."
    )

    add_heading(doc, "2.2 Core Components", level=2)
    add_bullet(doc, "User Application — Mobile/web interface for patients")
    add_bullet(doc, "Kiosk Website — Touchscreen terminal for community health centers")
    add_bullet(doc, "AI Caller System — Voice-based assessment for low-literacy users")
    add_bullet(doc, "Wearable Sensor Hardware — Smartwatch-style vital sign monitoring")
    add_bullet(doc, "Admin / NGO Monitoring Portal — Dashboard for health authorities")

    add_image_placeholder(doc, "component_diagram")

    add_page_break(doc)


def section_04_architecture(doc):
    """Page 4 — System Architecture"""
    add_heading(doc, "3. System Architecture")

    add_heading(doc, "3.1 Architecture Diagram", level=2)
    add_image_placeholder(doc, "architecture_diagram", width=Inches(5.5))

    add_heading(doc, "3.2 Data Flow Overview", level=2)
    add_body(doc,
        "Data flows from the user-facing interfaces (mobile app, kiosk, AI caller) to the "
        "backend AI engine. The backend processes responses through a RAG-based assessment "
        "pipeline, classifies severity, generates health reports, and stores results in a "
        "PostgreSQL database. The admin portal retrieves aggregated data for disease tracking "
        "and geographic analysis."
    )

    add_page_break(doc)


def section_05_user_workflow(doc):
    """Page 5 — User Workflow"""
    add_heading(doc, "4. User Workflow")

    add_heading(doc, "4.1 User Journey", level=2)
    add_image_placeholder(doc, "user_workflow_diagram")

    add_heading(doc, "4.2 Assessment Pipeline", level=2)
    add_body(doc,
        "The assessment pipeline uses a Retrieval-Augmented Generation (RAG) approach to "
        "dynamically generate follow-up questions based on user responses. Each completed "
        "assessment is processed through a severity classification engine that categorizes "
        "the health concern into actionable tiers."
    )
    add_bullet(doc, "RAG-based adaptive question generation")
    add_bullet(doc, "Multi-tier severity classification (Emergency / Doctor Visit / Self Care)")
    add_bullet(doc, "Automated report generation with recommendations")

    add_page_break(doc)


def section_06_mobile_app(doc):
    """Pages 6–7 — Mobile / User Application"""
    add_heading(doc, "5. User Application")

    add_heading(doc, "5.1 Application Purpose", level=2)
    add_body(doc,
        "The user application serves as the primary interface for patients to interact with "
        "the DeepBlue system. Available on mobile and web, it provides guided health assessments, "
        "report viewing, and AI-powered chat support."
    )

    add_heading(doc, "5.2 Key Features", level=2)
    add_bullet(doc, "User registration and profile management")
    add_bullet(doc, "Symptom entry and body region selection")
    add_bullet(doc, "AI-powered adaptive questionnaire")
    add_bullet(doc, "Health report generation and history")
    add_bullet(doc, "AI chatbot for health guidance")

    add_heading(doc, "5.3 Screenshots", level=2)
    add_image_placeholder(doc, "app_login")
    add_image_placeholder(doc, "app_questionnaire")
    add_image_placeholder(doc, "app_report")
    add_image_placeholder(doc, "app_chatbot")

    add_page_break(doc)


def section_07_kiosk(doc):
    """Pages 8–9 — Kiosk Website Interface"""
    add_heading(doc, "6. Kiosk Website Interface")

    add_heading(doc, "6.1 Purpose of Kiosk Deployment", level=2)
    add_body(doc,
        "Kiosk terminals enable health assessments in community health centers, pharmacies, "
        "and rural clinics. They provide a guided touchscreen experience for users who may not "
        "own smartphones, ensuring the system reaches the widest possible population."
    )

    add_heading(doc, "6.2 Interface Overview", level=2)
    add_bullet(doc, "Full assessment interface with guided navigation")
    add_bullet(doc, "Health report display and print capability")
    add_bullet(doc, "Integrated AI chatbot access")

    add_heading(doc, "6.3 Screenshots", level=2)
    add_image_placeholder(doc, "kiosk_assessment")
    add_image_placeholder(doc, "kiosk_report")
    add_image_placeholder(doc, "kiosk_chatbot")

    add_page_break(doc)


def section_08_ai_caller(doc):
    """Page 10 — AI Caller System"""
    add_heading(doc, "7. AI Caller System")

    add_heading(doc, "7.1 Voice Interaction System", level=2)
    add_body(doc,
        "The AI Caller system enables health assessments through voice interaction. The system "
        "calls the user, asks health-related questions in natural language, processes verbal "
        "responses, and conducts a complete assessment through conversation."
    )

    add_heading(doc, "7.2 Advantages", level=2)
    add_bullet(doc, "Accessibility for users without smartphones or internet access")
    add_bullet(doc, "Support for low-literacy populations")
    add_bullet(doc, "Hands-free, conversation-driven interaction")
    add_bullet(doc, "Works on basic feature phones")

    add_image_placeholder(doc, "ai_caller_workflow")

    add_page_break(doc)


def section_09_hardware(doc):
    """Pages 11–12 — Wearable Sensor Hardware"""
    add_heading(doc, "8. Wearable Sensor Hardware")

    add_heading(doc, "8.1 Hardware Overview", level=2)
    add_body(doc,
        "DeepBlue integrates a smartwatch-style wearable sensor module that captures real-time "
        "vital signs. The sensor data complements the questionnaire-based assessment, providing "
        "objective health metrics for more accurate severity classification."
    )

    add_heading(doc, "8.2 Sensor Components", level=2)
    add_bullet(doc, "Heart rate / pulse oximeter sensor")
    add_bullet(doc, "Body temperature sensor")
    add_bullet(doc, "Motion / accelerometer sensor")

    add_heading(doc, "8.3 Hardware Integration", level=2)
    add_body(doc,
        "Sensor data is transmitted via Bluetooth to the user application or kiosk terminal, "
        "where it is forwarded to the backend for analysis alongside questionnaire responses."
    )

    add_image_placeholder(doc, "hardware_prototype")
    add_image_placeholder(doc, "hardware_integration_diagram")

    add_page_break(doc)


def section_10_assessment_logic(doc):
    """Page 13 — Health Assessment Logic"""
    add_heading(doc, "9. Health Assessment Logic")

    add_heading(doc, "9.1 Severity Classification System", level=2)
    add_table(doc,
        ["Severity", "Color", "Meaning"],
        [
            ["Emergency", "Red", "Immediate medical attention required"],
            ["Doctor Visit", "Yellow", "Medical consultation recommended"],
            ["Self Care", "Green", "Home monitoring and self-care advised"],
        ],
    )

    add_heading(doc, "9.2 AI Decision Pipeline", level=2)
    add_body(doc,
        "The decision pipeline combines RAG-based contextual analysis, structured questionnaire "
        "scoring, and optional sensor data to produce a severity classification. The AI chatbot "
        "provides additional guidance based on the assessment outcome."
    )

    add_image_placeholder(doc, "ai_assessment_pipeline")

    add_page_break(doc)


def section_11_admin_portal(doc):
    """Pages 14–16 — Admin & NGO Monitoring Portal"""
    add_heading(doc, "10. Admin & NGO Monitoring Portal")

    add_heading(doc, "10.1 Purpose", level=2)
    add_body(doc,
        "The monitoring portal provides health authorities and NGOs with real-time visibility "
        "into disease patterns, severity distribution, and geographic health data. It enables "
        "coordinated response and resource allocation."
    )

    add_heading(doc, "10.2 Dashboard", level=2)
    add_body(doc,
        "The dashboard displays key performance indicators, active alerts, and disease "
        "distribution summaries."
    )
    add_image_placeholder(doc, "admin_dashboard")

    add_heading(doc, "10.3 Disease Map", level=2)
    add_body(doc,
        "An interactive geographic map displays disease reports with severity markers and "
        "filtering capabilities, enabling spatial analysis of health trends."
    )
    add_image_placeholder(doc, "admin_disease_map")

    add_heading(doc, "10.4 State-Level Analysis", level=2)
    add_body(doc,
        "State-level views provide aggregated insights including case statistics, NGO "
        "assignments, and regional health trends."
    )
    add_image_placeholder(doc, "admin_state_analysis")

    add_page_break(doc)


def section_12_analytics(doc):
    """Page 17 — Data & Analytics"""
    add_heading(doc, "11. Data & Analytics")

    add_heading(doc, "11.1 Disease Tracking", level=2)
    add_body(doc,
        "The system tracks disease types, frequency, and geographic spread across all "
        "deployment locations, building a comprehensive health dataset over time."
    )

    add_heading(doc, "11.2 Analytics Insights", level=2)
    add_bullet(doc, "Severity distribution across regions")
    add_bullet(doc, "State-level health ranking")
    add_bullet(doc, "Temporal trends in disease reporting")
    add_bullet(doc, "Assessment completion rates")

    add_page_break(doc)


def section_13_tech_stack(doc):
    """Page 18 — Technology Stack"""
    add_heading(doc, "12. Technology Stack")

    add_heading(doc, "Frontend", level=2)
    add_bullet(doc, "React + TypeScript")
    add_bullet(doc, "Tailwind CSS")
    add_bullet(doc, "Leaflet (interactive maps)")
    add_bullet(doc, "Kotlin Multiplatform (mobile)")

    add_heading(doc, "Backend", level=2)
    add_bullet(doc, "FastAPI (Python)")
    add_bullet(doc, "PostgreSQL / Supabase")
    add_bullet(doc, "AWS ECS (deployment)")

    add_heading(doc, "AI / ML", level=2)
    add_bullet(doc, "Google Gemini")
    add_bullet(doc, "RAG assessment pipeline")
    add_bullet(doc, "Gemini Vision (image analysis)")

    add_heading(doc, "Voice / Calling", level=2)
    add_bullet(doc, "Twilio")
    add_bullet(doc, "Cerebras (real-time voice AI)")

    add_heading(doc, "Hardware", level=2)
    add_bullet(doc, "ESP32 / Arduino-based sensor module")
    add_bullet(doc, "Bluetooth Low Energy (BLE)")

    add_page_break(doc)


def section_14_impact(doc):
    """Page 19 — Impact & Use Cases"""
    add_heading(doc, "13. Impact & Use Cases")

    add_body(doc,
        "DeepBlue addresses real-world healthcare gaps with tangible, deployable solutions:"
    )

    add_bullet(doc, "Rural Health Screening — Enables early detection in areas without doctors")
    add_bullet(doc, "NGO Coordination — Centralized disease monitoring for aid organizations")
    add_bullet(doc, "Outbreak Early Warning — Geographic clustering of cases signals potential outbreaks")
    add_bullet(doc, "Remote Healthcare Access — Voice-based and kiosk-based interfaces reach underserved populations")
    add_bullet(doc, "Data-Driven Policy — Aggregated health data supports public health decision-making")

    add_page_break(doc)


def section_15_future(doc):
    """Page 20 — Future Enhancements"""
    add_heading(doc, "14. Future Enhancements")

    add_bullet(doc, "Additional sensor types (blood pressure, ECG)")
    add_bullet(doc, "Hospital system integration for referral workflows")
    add_bullet(doc, "National-scale disease monitoring dashboard")
    add_bullet(doc, "Predictive outbreak modeling using ML")
    add_bullet(doc, "Multi-language support for regional accessibility")
    add_bullet(doc, "Offline-capable kiosk mode for areas without internet")

    add_page_break(doc)


def section_16_appendix(doc):
    """Page 21 — Appendix (Optional)"""
    add_heading(doc, "Appendix")

    add_body(doc, "Additional screenshots, diagrams, and technical references.")

    add_heading(doc, "A.1 Additional Screenshots", level=2)
    add_image_placeholder(doc, "appendix_screenshot_1")
    add_image_placeholder(doc, "appendix_screenshot_2")

    add_heading(doc, "A.2 Additional Diagrams", level=2)
    add_image_placeholder(doc, "appendix_diagram_1")


# ══════════════════════════════════════════════════════════════════════════
# MAIN — Assemble all sections
# ══════════════════════════════════════════════════════════════════════════

SECTIONS = [
    section_01_cover,
    section_02_overview,
    section_03_system_overview,
    section_04_architecture,
    section_05_user_workflow,
    section_06_mobile_app,
    section_07_kiosk,
    section_08_ai_caller,
    section_09_hardware,
    section_10_assessment_logic,
    section_11_admin_portal,
    section_12_analytics,
    section_13_tech_stack,
    section_14_impact,
    section_15_future,
    section_16_appendix,
]


def build():
    doc = Document()
    set_document_defaults(doc)

    for section_fn in SECTIONS:
        section_fn(doc)

    doc.save(str(OUTPUT_PATH))
    print(f"✔ Dossier saved → {OUTPUT_PATH}")
    print(f"  Images dir   → {IMAGES_DIR}/")
    print(f"  Sections     → {len(SECTIONS)}")
    print()
    print("To add images, drop files into docs/images/ with these names:")
    print("  (png/jpg accepted)")
    for name in [
        "cover_illustration", "component_diagram", "architecture_diagram",
        "user_workflow_diagram", "app_login", "app_questionnaire", "app_report",
        "app_chatbot", "kiosk_assessment", "kiosk_report", "kiosk_chatbot",
        "ai_caller_workflow", "hardware_prototype", "hardware_integration_diagram",
        "ai_assessment_pipeline", "admin_dashboard", "admin_disease_map",
        "admin_state_analysis", "appendix_screenshot_1", "appendix_screenshot_2",
        "appendix_diagram_1",
    ]:
        print(f"    {name}.png")


if __name__ == "__main__":
    build()
